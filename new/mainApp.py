#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from PyQt5 import QtCore, QtWidgets, QtMultimedia
from docx import Document, shared
from shell_ui import Ui_MainWindow
from CustomDateEdit import DateEdit as customDateEdit
from dialog import MyDialog
from exception_handler import catch_exceptions
import os
import shutil
import json
import sys



class MyWin(QtWidgets.QMainWindow, Ui_MainWindow):

    def __init__(self, parent=None):
        super(MyWin, self).__init__(parent)
        #super().__init__()
        #QtWidgets.QWidget.__init__(self, parent)
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)

        self.ui.gridLayout.removeWidget(self.ui.dateEdit)
        self.ui.dateEdit.close()
        self.ui.dateEdit = customDateEdit(self.ui.centralwidget)
        self.ui.gridLayout.addWidget(self.ui.dateEdit, 3, 1, 1, 2)
        self.ui.gridLayout.update()

        #self.ui.gridLayout.addWidget(customDateEdit, 3, 1, 1, 2)

        self.ui.lineEdit_4.setText(os.getcwd())

        self.DOCTORS=[] 
        self.VESSELS=[] 
        self.PROTOCOL_ULTRASOUND=[]
        self.typeOfProtocols={
            'Приём Врача':'Доктора',
            'УЗИ Общее':'Протоколы УЗИ',
            'УЗИ Сосудов':'Сосуды',
        }



        self.ui.pushButton.clicked.connect(self.close)
        self.ui.pushButton_2.clicked.connect(self.copyAndOpenTheTemplate)
        self.ui.lineEdit.textChanged.connect(self.enableButton)
        self.ui.lineEdit_2.textChanged.connect(self.enableButton)
        self.ui.lineEdit_3.textChanged.connect(self.enableButton)
        #self.ui.menu.addAction("Выход", self.close)
        #self.ui.menu.addAction("Выбор папок", self.menuChoice)
        self.ui.action.triggered.connect(self.menuChoice)
        self.ui.action_2.triggered.connect(self.close)
        self.ui.pushButton_3.clicked.connect(self.fillBirthDate)
        self.ui.buttonGroup.buttonToggled.connect(self.enableButton)
        self.ui.buttonGroup.buttonClicked.connect(self.fullfillComboBox)
        self.ui.toolButton.clicked.connect(self.selectFolder)
        self.ui.pushButton_4.clicked.connect(self.rmvFromSrv)
        self.ui.pushButton_5.clicked.connect(self.addToSrv)
        self.ui.listWidget.doubleClicked.connect(self.addToSrv)
        


        self.readSettingsFile()


                    #self.ui.comboBox.addItem(str(_))
                    #print(_)

    def copy_rename(self, old_file_name, new_file_name, src_dir_file, dst_dir):
        try:
            #src_dir=os.path.join(os.curdir,'Протоколы УЗИ')
            #dst_dir= dest_dir #os.path.join(os.curdir , new_folder_name)
            if (os.path.isdir(dst_dir)):
                print('folder already exist')
            else:
                os.mkdir(dst_dir)
            #print('dst_dir',dst_dir)
            #print('src_dir_file',src_dir_file)
            shutil.copy(src_dir_file,dst_dir)
            #print('src_file',src_dir_file)
            #print('---->>>',os.listdir(dst_dir))
            
            dst_file = os.path.join(dst_dir, old_file_name)
            new_dst_file_name = os.path.join(dst_dir, new_file_name)
            os.rename(dst_file, new_dst_file_name)

            # add with name and surname
            document = Document(new_dst_file_name)

            paragraphs = document.paragraphs
            paragraphs[1].style.font.size = shared.Pt(12)
            text = paragraphs[1].text
            # patient = dst_dir.split('_')[0]+' '+dst_dir.split('_')[1]+' '+dst_dir.split('_')[2]+'\n'\
            #     +dst_dir.split('_')[3]+'.'+dst_dir.split('_')[4]+'.'+dst_dir.split('_')[5]+'\n'
            patient = 'Ф.И.О.: '  + self.ui.lineEdit.text() + ' ' + self.ui.lineEdit_2.text() + ' ' + self.ui.lineEdit_3.text()+'\n'\
                + 'Дата рождения: ' + self.ui.dateEdit.date().toString('dd.MM.yyyy') + ' ' \
                + 'Дата исследования: ' + QtCore.QDate.currentDate().toString('dd.MM.yyyy') + '\n'
            paragraphs[1]._p.clear()
            paragraphs[1].add_run(patient + text)
            document.save(new_dst_file_name)
        except OSError:
            print('OSError')
    
    def enableButton(self):
        # if self.ui.lineEdit.text().isEmpty(): # print('Empty') 
        if (self.ui.lineEdit.text() != "" and self.ui.lineEdit_2.text() != "" and self.ui.lineEdit_3.text() != "" \
            and self.ui.buttonGroup.checkedButton() is not None):
            self.ui.pushButton_2.setEnabled(True)
        else:
            self.ui.pushButton_2.setEnabled(False)
        #print(self.ui.buttonGroup.checkedButton())
    
    def menuChoice(self):
        dialog = MyDialog(self)
        result = dialog.exec_()
        if result == QtWidgets.QDialog.Accepted:
            print("Нажата кнопка Save")
            # self.ui.lineEdit_4.setText(dialog.dialog_ui.lineEdit_4.text())
            self.readSettingsFile()
        else:
            print("Нажата кнопка Cancel, кнопка Закрыть или клавиша <Esc>", result)
        #print('ssssss')

    # def closeProgram(self):
    #     self.close()
    
    def fullfillComboBox(self):
        # print(self.ui.buttonGroup.checkedButton().text())
        selectButton=self.ui.buttonGroup.checkedButton().text()
        if selectButton=='Приём Врача':
            # self.ui.comboBox.addItems(self.DOCTORS)
            self.ui.listWidget.clear()
            self.ui.listWidget.addItems(self.DOCTORS)
            print('DOCTORS')
        elif selectButton=='УЗИ Общее':
            # self.ui.comboBox.addItems(self.PROTOCOL_ULTRASOUND)
            self.ui.listWidget.clear()
            self.ui.listWidget.addItems(self.PROTOCOL_ULTRASOUND)
            print('PROTOCOL_ULTRASOUND')
        elif selectButton=='УЗИ Сосудов':
            # self.ui.comboBox.addItems(self.VESSELS)
            self.ui.listWidget.clear()
            self.ui.listWidget.addItems(self.VESSELS)
            print('VESSELS')

    def selectFolder(self):
        directory=QtWidgets.QFileDialog.getExistingDirectory(self, "Select Directory")
        if (directory):
            self.ui.lineEdit_4.setText(directory)
            with open("settings.json", "r") as jsonFile:
                data = json.load(jsonFile)
            data["Пациенты"] = directory
            with open("settings.json", "w") as jsonFile:
                json.dump(data, jsonFile)
            #print('No')
        #print(directory)
    
    def readSettingsFile(self):
        self.DOCTORS.clear()
        self.PROTOCOL_ULTRASOUND.clear()
        self.VESSELS.clear()
        with open('settings.json', 'r') as jsonfile:
            self.jsonData = json.load(jsonfile)

            self.ui.lineEdit_4.setText(self.jsonData['Пациенты'])

            try:
                completer = QtWidgets.QCompleter( [client.split('_')[0] for client in os.listdir(self.jsonData['Пациенты'])], self.ui.lineEdit)
                self.ui.lineEdit.setCompleter(completer)
                completer = QtWidgets.QCompleter( [client.split('_')[1] for client in os.listdir(self.jsonData['Пациенты'])], self.ui.lineEdit_2)
                self.ui.lineEdit_2.setCompleter(completer)
                completer = QtWidgets.QCompleter( [client.split('_')[2] for client in os.listdir(self.jsonData['Пациенты'])], self.ui.lineEdit_3)
                self.ui.lineEdit_3.setCompleter(completer)
            except:
                print('Пациентов не было найдено')

            for _ in os.listdir(self.jsonData['Доктора']):
                self.DOCTORS.append(_.split('.')[0])
            for _ in os.listdir(self.jsonData['Протоколы УЗИ']):
                self.PROTOCOL_ULTRASOUND.append(_.split('.')[0])
            for _ in os.listdir(self.jsonData['Сосуды']):
                self.VESSELS.append(_.split('.')[0])
            self.ui.lineEdit.setText(self.jsonData['Фамилия'])
            self.ui.lineEdit_2.setText(self.jsonData['Имя'])
            self.ui.lineEdit_3.setText(self.jsonData['Отчество'])
            # birthDateStr = jsonData['ДатаРождения'] #"20/12/2015";
            # QDate Date = QDate::fromString(date_string_on_db,"dd/MM/yyyy");
            self.ui.dateEdit.setDate(QtCore.QDate.fromString(self.jsonData['ДатаРождения'],"dd.MM.yyyy")) # (1993, 12, 25)  # .setText(jsonData['Отчетсво'])
    
    def fillBirthDate(self):
        try:
            for client in os.listdir(self.jsonData['Пациенты']):
                if client.split('_')[0] == self.ui.lineEdit.text() and \
                   client.split('_')[1] == self.ui.lineEdit_2.text() and \
                   client.split('_')[2] == self.ui.lineEdit_3.text():
                    # print(str(client.split('_')[3:]))
                    self.ui.dateEdit.setDate(QtCore.QDate.fromString('_'.join(client.split('_')[3:]), "dd_MM_yyyy"))
                    self.ui.pushButton_3.setEnabled(False)
        except:
            print('Не найдено года рождения')
        


    def copyAndOpenTheTemplate(self):
        #copy_rename(old_file_name, new_file_name, src_dir, dest_dir):

        now = QtCore.QDate.currentDate().toString('dd_MM_yyyy') # datetime.date(2012, 12, 14).strftime('%d_%m_%Y'))
        birthDate=self.ui.dateEdit.date().toString('dd_MM_yyyy')
        name=self.ui.lineEdit_2.text()
        surname=self.ui.lineEdit.text()
        middleName=self.ui.lineEdit_3.text()

        # protocol=self.ui.comboBox.currentText()
        # with open('settings.json', 'r') as jsonfile:
        dirWithProtocols=self.jsonData[self.typeOfProtocols[self.ui.buttonGroup.checkedButton().text()]]
            
 
        # protocol = [filename for filename in os.listdir(dirWithProtocols) if filename.startswith(self.ui.comboBox.currentText())][0]
        # print('--->',protocol)
        for item in [self.ui.listWidget_2.item(i) for i in range(  self.ui.listWidget_2.count())]:
            protocolName=item.text()
            protocolType=self.typeOfProtocols[item.data(QtCore.Qt.UserRole)]

            dirWithProtocols=self.jsonData[self.typeOfProtocols[item.data(QtCore.Qt.UserRole)]]
            protocol = [filename for filename in os.listdir(dirWithProtocols) if filename.startswith(item.text())][0]
            protocolExt=protocol.split('.')[1] #self.ui.comboBox.currentText().split('.')[1]

            # protocol=item.text() + '.' + protocolExt
            newDirName=('{0}_{1}_{2}_{3}').format(surname,name,middleName,birthDate)
            newFileName=('{0}_{1}.{2}').format(protocolName,now,protocolExt)
            destDir=os.path.join(self.ui.lineEdit_4.text(),newDirName)
            srcDir=os.path.join(os.getcwd(),protocolType,protocol)# self.ui.comboBox.currentText())
            new_dst_file_name=os.path.join(destDir,newFileName)
            self.copy_rename(protocol, newFileName, srcDir, destDir)
    
            command='start winword "{}"'.format(new_dst_file_name.replace('/', '\\'))
            os.system(command)

        # protocolName=protocol.split('.')[0]
        # protocolExt='docx' #protocol.split('.')[1] #self.ui.comboBox.currentText().split('.')[1]
        # newDirName=('{0}_{1}_{2}_{3}').format(surname,name,middleName,birthDate)
        # newFileName=('{0}_{1}.{2}').format(protocolName,now,protocolExt)
        # destDir=os.path.join(self.ui.lineEdit_4.text(),newDirName)
        # srcDir=os.path.join(os.getcwd(),self.typeOfProtocols[self.ui.buttonGroup.checkedButton().text()],protocol)# self.ui.comboBox.currentText())
        # new_dst_file_name=os.path.join(destDir,newFileName)
        # print(new_dst_file_name)
        # self.copy_rename(protocol, newFileName, srcDir, destDir)
        # os.system(command)

        self.jsonData['Фамилия'] = surname 
        self.jsonData['Имя'] = name
        self.jsonData['Отчество'] = middleName
        self.jsonData['ДатаРождения'] =  self.ui.dateEdit.date().toString('dd.MM.yyyy')
        # settings = json.dumps(settings).decode('unicode-escape').encode('utf8')
        with open('settings.json', 'w') as json_file:  
            json.dump(self.jsonData, json_file, ensure_ascii=False)
        print('OKAY!')
        self.close()
    
    def addToSrv(self):
        if not (self.ui.listWidget.selectedIndexes() == []):
            item=QtWidgets.QListWidgetItem(self.ui.listWidget.currentItem().text())
            item.setData(QtCore.Qt.UserRole, self.ui.buttonGroup.checkedButton().text())
            #   data = currentItem->data(Qt::UserRole)
            selectedValues = [self.ui.listWidget_2.item(i).text() for i in range(self.ui.listWidget_2.count())]
            if item.text() not in selectedValues:
                self.ui.listWidget_2.addItem(item)#self.ui.listWidget.currentItem().text())
            # self.enableButton()

    def rmvFromSrv(self):
        listItems=self.ui.listWidget_2.selectedItems()
        if not listItems: return        
        for item in listItems:
            print(item.text())
            print(item.data(QtCore.Qt.UserRole))
            self.ui.listWidget_2.takeItem(self.ui.listWidget_2.row(item))
        print('RemovevFromSRV')

 
if __name__ == "__main__":
    if not ('settings.json' in os.listdir()):
        settings = {  
            'Доктора': os.path.join(os.getcwd(),'Доктора'),
            'Протоколы УЗИ': os.path.join(os.getcwd(),'Протоколы УЗИ'),
            'Сосуды': os.path.join(os.getcwd(),'Сосуды'),
            'Пациенты': os.getcwd(),
            'Фамилия': "Иванов",
            'Имя': "Иван",
            'Отчество': "Иванович",
            'ДатаРождения': "25.12.1993"
        }
        # settings = json.dumps(settings).decode('unicode-escape').encode('utf8')
        with open('settings.json', 'w') as json_file:  
            json.dump(settings, json_file, ensure_ascii=False)
    app = QtWidgets.QApplication(sys.argv)
    #ico = QtWidgets.QWidget().style().standardIcon(QtWidgets.QStyle.SP_MediaVolume)
    #app.setWindowIcon(ico)
    w = MyWin()
    w.show()
sys.exit(app.exec())